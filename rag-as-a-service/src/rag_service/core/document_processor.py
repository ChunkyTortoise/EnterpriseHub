"""Document processor — parse PDF/CSV/TXT/MD/DOCX and chunk text."""

from __future__ import annotations

import csv
import io
import logging
import re
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class TextChunk:
    """A chunk of text from a document."""
    content: str
    chunk_index: int
    metadata: dict = field(default_factory=dict)


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    filename: str
    content_type: str
    raw_text: str
    chunks: list[TextChunk]
    metadata: dict = field(default_factory=dict)


class DocumentProcessor:
    """Parse documents and split into chunks for embedding."""

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "text/csv": "csv",
        "text/plain": "txt",
        "text/markdown": "md",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    }

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def process(self, content: bytes, filename: str, content_type: str) -> ProcessedDocument:
        """Process a document and return chunked text."""
        doc_type = self.SUPPORTED_TYPES.get(content_type)
        if doc_type is None:
            # Try to infer from filename extension
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
            type_map = {"pdf": "pdf", "csv": "csv", "txt": "txt", "md": "md", "docx": "docx"}
            doc_type = type_map.get(ext)
            if doc_type is None:
                raise ValueError(f"Unsupported content type: {content_type}")

        raw_text = self._extract_text(content, doc_type)
        chunks = self._chunk_text(raw_text, filename)

        return ProcessedDocument(
            filename=filename,
            content_type=content_type,
            raw_text=raw_text,
            chunks=chunks,
            metadata={"doc_type": doc_type, "char_count": len(raw_text)},
        )

    def _extract_text(self, content: bytes, doc_type: str) -> str:
        """Extract text based on document type."""
        extractors = {
            "pdf": self._extract_pdf,
            "csv": self._extract_csv,
            "txt": self._extract_text_file,
            "md": self._extract_text_file,
            "docx": self._extract_docx,
        }
        return extractors[doc_type](content)

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyPDF2 if available, else basic fallback."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            # Fallback: try to extract text between stream markers
            text = content.decode("latin-1", errors="ignore")
            # Remove binary content
            text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", text)
            return text.strip()

    def _extract_csv(self, content: bytes) -> str:
        """Extract text from CSV, converting rows to readable text."""
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ""

        headers = rows[0] if rows else []
        lines = []
        for row in rows[1:]:
            parts = [f"{h}: {v}" for h, v in zip(headers, row) if v.strip()]
            lines.append("; ".join(parts))

        return "\n".join(lines)

    def _extract_text_file(self, content: bytes) -> str:
        """Extract text from plain text or markdown."""
        return content.decode("utf-8", errors="replace")

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx if available."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            # Basic fallback: DOCX is a zip with XML
            import zipfile

            try:
                with zipfile.ZipFile(io.BytesIO(content)) as zf:
                    if "word/document.xml" in zf.namelist():
                        xml = zf.read("word/document.xml").decode("utf-8")
                        # Strip XML tags
                        text = re.sub(r"<[^>]+>", " ", xml)
                        return re.sub(r"\s+", " ", text).strip()
            except zipfile.BadZipFile:
                pass
            return ""

    def _chunk_text(self, text: str, filename: str) -> list[TextChunk]:
        """Split text into overlapping chunks by character count."""
        if not text.strip():
            return []

        # Approximate token-to-char ratio (1 token ~ 4 chars)
        char_size = self.chunk_size * 4
        char_overlap = self.chunk_overlap * 4

        chunks: list[TextChunk] = []
        start = 0
        idx = 0

        while start < len(text):
            end = start + char_size
            chunk_text = text[start:end].strip()

            if chunk_text:
                chunks.append(
                    TextChunk(
                        content=chunk_text,
                        chunk_index=idx,
                        metadata={"source": filename, "start_char": start, "end_char": end},
                    )
                )
                idx += 1

            if end >= len(text):
                break

            start = end - char_overlap

        return chunks
